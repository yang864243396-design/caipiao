# -*- coding: utf-8 -*-
"""生成面向客户的项目进度汇报 Word 文档（口语化版本）。"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent / "项目进度汇报-客户版.docx"
REPORT_DATE = date(2026, 6, 15)


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    sizes = {1: 16, 2: 14, 3: 12}
    set_run_font(run, size=sizes.get(level, 11), bold=True, color=(0, 80, 203) if level == 1 else None)
    return p


def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, size=10, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, size=10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # 封面
    for _ in range(4):
        doc.add_paragraph()
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover_title.add_run("彩票方案平台\n开发进度说明")
    set_run_font(r, size=22, bold=True, color=(0, 80, 203))

    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = cover_sub.add_run(f"{REPORT_DATE.strftime('%Y年%m月%d日')}")
    set_run_font(r2, size=12)

    doc.add_page_break()

    # 先说清楚这个项目是啥
    add_title(doc, "一、这个项目是干啥的", 1)
    add_body(
        doc,
        "简单说，就是给你们做一套完整的彩票方案平台：会员在网页上登录、看彩种、配方案、"
        "挂到云端自动跑；运营人员在后台管会员、管方案、看数据、发公告。",
        indent=True,
    )
    add_body(
        doc,
        "技术上是三块：会员用的网页（用户端）、运营用的后台（管理端）、以及跑业务的 Go 后端 + PostgreSQL 数据库。"
        "另外还做了消息推送，比如维护开关、余额变动、开奖结果这些不用手动刷新页面。",
        indent=True,
    )
    add_body(
        doc,
        "目前大框架和主要功能都写完了，本地/内测环境可以完整跑一遍演示。"
        "还在搞的是：跟第三方挂机平台的深度对接，以及几种高级方案玩法的前端界面。",
        indent=True,
    )

    add_title(doc, "1.1 做了哪些东西", 2)
    add_table(
        doc,
        ["东西", "大概是什么", "现在啥情况"],
        [
            ["会员网页", "大厅、配方案、云端跑方案、看投注、会员中心、公告聊天等", "基本做完了"],
            ["运营后台", "看板、查会员、管方案、管订单、发内容、管账号权限等", "基本做完了"],
            ["后端服务", "所有接口、登录验证、业务逻辑、后台自动任务", "基本做完了"],
            ["数据库", "表结构、演示数据、升级脚本", "做完了"],
            ["接口文档", "前后端怎么对接、怎么验收，都写好了", "做完了"],
            ["第三方挂机对接", "绑号、看真实余额、下单派奖走第三方", "正在做，大概四成"],
            ["高级方案类型", "7 种不同出号方式，前后端都要改", "后端好了，前端还在改"],
        ],
        col_widths=[3.0, 7.5, 3.0],
    )

    add_title(doc, "1.2 整体进度（大概估的）", 2)
    add_table(
        doc,
        ["看哪块", "进度", "一句话说明"],
        [
            ["平台主体（能内测演示）", "约 85%", "登录到大厅到跑方案这条主链路通了"],
            ["运营后台", "约 98%", "就客服聊天还占着位，别的都能用"],
            ["第三方挂机对接", "约 40%", "架子搭好了，还在联调"],
            ["高级方案引擎", "约 60%", "后台逻辑写完了，页面还在改"],
            ["能正式上线", "约 75%～80%", "等第三方对接跑通 + 上线前几个事定下来"],
        ],
        col_widths=[4.0, 2.0, 7.5],
    )

    doc.add_page_break()

    # 已经做完的
    add_title(doc, "二、已经做完的主要工作", 1)
    add_body(
        doc,
        "下面按模块列一下，方便你们对照。每一块都是需求梳理、写接口、写页面、建表、联调一起搞下来的。",
        indent=True,
    )

    modules = [
        (
            "2.1 基础能力（登录、维护这些）",
            [
                "会员端和后台都能登录，token 验证通了",
                "全站维护开关：后台一开，前台立刻进维护页（不用刷新）",
                "接口报错统一处理，环境配置模板也准备好了",
                "服务健康检查、数据库连接检测，启动脚本都 OK",
            ],
        ),
        (
            "2.2 会员网页（36 个页面）",
            [
                "游戏大厅：看彩种、维护中的进不去",
                "跟单大厅：看排行榜、点进去玩游戏",
                "方案这块：从分享池下载、自己建方案、改倍投和期次、跟单/反买",
                "云端中心：看正在跑的方案，能开、停、暂停、恢复，还能设「回头」规则",
                "投注记录：真实模式和模拟模式分开看，有汇总也有明细",
                "游戏详情：看期号倒计时、自己选手动下注、查历史开奖",
                "会员中心：资料、余额、投注记录、流水、追号、充提记录（充值目前是演示用的，点完就到账）",
                "内容互动：公告、常见问题、帮助、意见反馈、聊天、系统消息",
                "第三方账号绑定：绑号、切换账号的页面有了，门禁逻辑也接了（还在联调）",
            ],
        ),
        (
            "2.3 运营后台（24 个页面）",
            [
                "仪表盘：充值、提现、投注这些数字，有变化会自动刷新",
                "会员管理：搜索、看详情、查流水、做运营操作",
                "方案监控：看全站谁在跑方案，能强制停、解封，管分享池里的方案",
                "游戏运营：47 个彩种的上架/维护、跟单榜怎么排、方案模板库",
                "订单财务：查投注和追号、看帐变、审提现（演示链路）",
                "内容管理：公告、FAQ、帮助、大厅广告位、网站名称/logo、系统消息模板和下发",
                "系统管理：后台账号增删改、角色权限（不同角色看不同菜单）、操作日志、维护开关",
                "报表：按彩种统计、盈亏报表",
            ],
        ),
        (
            "2.4 后端和数据（比较重的一块）",
            [
                "大概 91 个接口，会员端、后台、公开接口都覆盖了",
                "5 种实时推送：维护、钱包变动、方案状态、聊天消息、新开奖",
                "85 份数据库升级脚本，30 多张业务表，演示数据也 seed 好了",
                "后台自动任务：方案自己跑、投注自动结算、开奖同步",
                "玩法结算逻辑：时时彩、PK10、快3、六合彩、PC28、11选5 都写了",
                "38 份自动化测试 + 一键冒烟脚本，改代码不容易改挂",
            ],
        ),
        (
            "2.5 彩种和玩法",
            [
                "47 个彩种、340 种玩法，数据都进库了",
                "旧的 9 个彩种有清理方案，首次部署会自动 purge 掉",
                "彩种可以设维护、改展示名、配第三方对接编码",
                "玩法树两层结构，下注前的校验规则也对齐了",
            ],
        ),
        (
            "2.6 文档",
            [
                "分阶段开发计划，前后端每个页面对应哪个接口，都列清楚了",
                "三端联调验收清单，150 多条，照着点就能验",
                "第三方挂机对接方案（产品那边已经定稿了）",
                "彩种迁移、方案模块、WebSocket 等技术说明",
                "OpenAPI 接口文档 4000 多行，前后端对着这个开发",
            ],
        ),
    ]

    for title, items in modules:
        add_title(doc, title, 2)
        add_bullets(doc, items)

    doc.add_page_break()

    # 工作量
    add_title(doc, "三、工作量大概多少", 1)
    add_body(
        doc,
        "用数字给你们感受一下投入规模（截至今天统计的）：",
        indent=True,
    )
    add_table(
        doc,
        ["啥", "多少", "备注"],
        [
            ["代码总量", "约 7.9 万行", "后端+前端+SQL+文档都算上了"],
            ["后端 Go 代码", "约 3.3 万行", "270 多个文件"],
            ["会员网页", "约 2.4 万行", "Vue3 + TypeScript"],
            ["运营后台", "约 6700 行", "Vue3 + TypeScript"],
            ["数据库脚本", "85 份", "建表、改表、演示数据"],
            ["接口", "约 91 个", "REST 的"],
            ["实时推送", "5 种", "断了还能轮询兜底"],
            ["页面", "60 个", "会员 36 + 后台 24"],
            ["自动化测试", "38 份", ""],
            ["技术文档", "16 篇", ""],
        ],
        col_widths=[3.5, 3.5, 5.5],
    )

    add_title(doc, "3.1 分阶段进度", 2)
    add_table(
        doc,
        ["阶段", "干啥的", "进度"],
        [
            ["第 0 阶段", "登录、维护开关这些基础", "100%"],
            ["第 1 阶段", "云端中心、投注记录、方案自动跑", "100%"],
            ["第 2 阶段", "会员钱包、订单、充提演示、团队推广", "100%"],
            ["第 3 阶段", "方案玩法、跟单大厅、游戏页下注", "100%"],
            ["第 4 阶段", "公告 FAQ 帮助、聊天、系统消息", "100%"],
            ["第 5 阶段", "运营后台全套", "98%"],
            ["专项：47 彩种", "新彩种玩法目录迁移", "约 88%"],
            ["专项：第三方对接", "Guaji 挂机平台", "约 40%"],
            ["专项：方案类型", "7 种不同出号方式", "约 60%"],
        ],
        col_widths=[2.5, 8.0, 2.0],
    )

    doc.add_page_break()

    # 还在做啥
    add_title(doc, "四、还在做啥、后面怎么排", 1)

    add_title(doc, "4.1 现在正在搞的", 2)
    add_table(
        doc,
        ["事项", "做到哪了", "为啥重要"],
        [
            [
                "第三方挂机平台对接",
                "后端绑号、切号、刷新 token、同步余额的代码有了；"
                "真实下单和派奖还在跟对方环境联调",
                "不上线这个，就没法用真实钱跑",
            ],
            [
                "7 种方案类型的前端界面",
                "后台跑方案的引擎写完了，测试也过了；"
                "会员配方案那几个页面还在改",
                "不然高级玩法会员配不了",
            ],
            [
                "对接后的页面裁剪",
                "按之前定的方案，充提、团队、银行卡、开户这些入口要拿掉",
                "钱走第三方了，这边界面得跟产品口径一致",
            ],
        ],
        col_widths=[3.5, 6.5, 3.5],
    )

    add_title(doc, "4.2 上线前还得定的事", 2)
    add_bullets(
        doc,
        [
            "后台客服聊天：现在是个占位页，要不要接、怎么接，得你们定",
            "系统消息：现在只能给单个会员发，要不要分群发、怎么分，也得定",
            "充值：现在是演示用的，点完就到账；正式上线得换真实支付",
            "一些运营开关：比如「一键关会员端入口」这种，产品还没最终拍板",
        ],
    )

    add_title(doc, "4.3 建议后面这么排", 2)
    add_table(
        doc,
        ["优先级", "干啥", "干完能怎样"],
        [
            ["最急", "第三方挂机联调跑通", "真实资金、下单、派奖整条链路能用"],
            ["最急", "该删的页面删干净，前后端对齐", "跟最终产品一致"],
            ["其次", "7 种方案类型页面验收", "会员能配、能跑"],
            ["其次", "按验收清单全跑一遍", "内测/预发可以签字"],
            ["后面", "部署上线、配域名证书", "能对外"],
            ["后面", "上线备忘里那些运营细节定稿", "客服、消息策略等"],
        ],
        col_widths=[1.8, 5.5, 6.2],
    )

    doc.add_page_break()

    # 需要客户配合的
    add_title(doc, "五、需要你们配合的地方", 1)
    add_bullets(
        doc,
        [
            "第三方挂机那边：联调进度很大程度看对方测试环境稳不稳、账号给不给力，"
            "麻烦帮忙协调一下测试账号和环境。",
            "对接完成后，本平台自己的充提、团队、银行卡这些会下线，"
            "用户充值提现得去第三方——文案和引导语麻烦你们确认一下怎么说。",
            "47 个彩种的玩法树建好之后就不能随便改了，后面要加玩法得发版，"
            "有运营需求的话尽早提。",
            "上线前最好一起跑 1～2 轮验收（我们有清单），产品和运营一起点一遍，签字确认。",
        ],
    )

    # 总结
    add_title(doc, "六、总结", 1)
    add_body(
        doc,
        "总的来说，这个平台从 0 搭到现在，会员端、后台、后端、数据库、实时推送都齐了，"
        "核心业务在内测环境能完整演示。大头已经干完了。",
        indent=True,
    )
    add_body(
        doc,
        "剩下主要是：第三方挂机深度对接、几种高级方案的前端界面、以及上线前几个运营细节定下来。"
        "我们会按优先级继续推，有进展随时同步。里程碑或者验收标准要调整的话，直接跟我们说就行。",
        indent=True,
    )

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sr = sig.add_run("开发组")
    set_run_font(sr, size=11)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(f"Generated: {path}")
    # 同步一份英文名，方便发文件
    import shutil
    shutil.copy2(path, path.parent / "project-progress-report-client.docx")
